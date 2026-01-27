from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def audit_to_docx(audit: dict) -> bytes:
    doc = Document()
    doc.add_heading('Marketing Kit Audit Report', 0)

    # Expected vs Rendered Sections
    doc.add_heading('Section Summary', level=1)
    doc.add_paragraph(f"Expected Sections: {', '.join(audit.get('expected_sections', []))}")
    doc.add_paragraph(f"Rendered Sections: {', '.join(audit.get('rendered_sections', []))}")
    missing = audit.get('missing_sections', [])
    if missing:
        doc.add_paragraph(f"Missing Sections: {', '.join(missing)}")
    else:
        doc.add_paragraph("No missing sections.")

    # Enhanced Section-by-Section Audit
    section_audits = audit.get('section_audits', [])
    for section in section_audits:
        doc.add_heading(f"Section: {section.get('section', '')}", level=2)
        status = section.get('status', '')
        doc.add_paragraph(f"Status: {status}")
        doc.add_paragraph(f"Reason: {section.get('reason', '')}")
        # Highlight missing findings in Key Findings section
        if section.get('section', '') == 'key_findings' and status == 'missing_findings':
            p = doc.add_paragraph()
            run = p.add_run("WARNING: No findings detected in Key Findings section!")
            run.bold = True
            run.font.color.rgb = (255, 0, 0)
        blocks = section.get('blocks', [])
        required_block_types = section.get('required_block_types', [])
        found_types = set(b.get('type') for b in blocks)
        missing_types = [t for t in required_block_types if t not in found_types]
        if missing_types:
            doc.add_paragraph(f"Missing block types: {', '.join(missing_types)}")
        # Check for empty text/fields and formatting issues
        issues_found = False
        if blocks:
            doc.add_heading('Block Details', level=3)
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Index'
            hdr_cells[1].text = 'Type'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Issues'
            hdr_cells[4].text = 'Text'
            hdr_cells[5].text = 'Formatting'
            for b in blocks:
                block_issues = list(b.get('issues', []))
                # Check for empty text
                if b.get('type') in ('Paragraph', 'Subhead', 'Callout') and not b.get('text', '').strip():
                    block_issues.append('Empty text')
                # Check for missing items in lists
                if b.get('type') in ('Bullets', 'Checklist', 'NumberedList') and not b.get('items'):
                    block_issues.append('No items')
                # Check for missing columns/rows in tables
                if b.get('type') == 'Table':
                    if not b.get('columns'):
                        block_issues.append('No columns')
                    if not b.get('rows'):
                        block_issues.append('No rows')
                # Formatting check (example: check for bold in Subhead)
                formatting = ''
                if b.get('type') == 'Subhead' and not b.get('text', '').isupper():
                    formatting = 'Not all caps'
                row_cells = table.add_row().cells
                row_cells[0].text = str(b.get('index', ''))
                row_cells[1].text = b.get('type', '')
                row_cells[2].text = b.get('status', '')
                row_cells[3].text = ", ".join(block_issues)
                row_cells[4].text = b.get('text', '')[:100]  # Truncate for readability
                row_cells[5].text = formatting
                if block_issues or formatting:
                    issues_found = True
        else:
            doc.add_paragraph("No blocks found in this section.")
        if not issues_found and not missing_types:
            doc.add_paragraph("All blocks present and complete.")

    # Save to bytes
    from io import BytesIO
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()
